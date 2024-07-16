from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from deep_translator import GoogleTranslator
import textwrap
import os
from docx import Document

def translate_pdf(pdf_path, dest_lang='en'):
    translator = GoogleTranslator(target=dest_lang)

    # Open the original PDF file
    reader = PdfReader(pdf_path)
    num_pages = len(reader.pages)
    
    # Get the base name of the PDF file without extension
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    translated_pdf_path = f"data/{pdf_name}.pdf"  # Adjust the path format as needed

    # Create a new canvas for the translated PDF
    c = canvas.Canvas(translated_pdf_path, pagesize=letter)

    # Register the TTF font
    pdfmetrics.registerFont(TTFont('Amiri', 'Amiri-Regular.ttf'))
    c.setFont("Amiri", 12)

    # Iterate through each page
    for page_num in range(num_pages):
        page = reader.pages[page_num]
        text = page.extract_text()

        # Split text into chunks of 4000 characters (adjust as needed)
        chunk_size = 4000
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

        # Translate each chunk and add to the translated PDF
        for chunk in chunks:
            translated_text = translator.translate(chunk)

            # Handle multi-line text properly
            lines = textwrap.wrap(translated_text, width=80)  # Adjust the width as needed
            y_position = 750  # Starting Y position for the text

            for line in lines:
                c.drawString(40, y_position, line)
                y_position -= 15  # Move down for the next line, adjust as needed

            # Add a new page if it's not the last chunk of the page
            if chunks.index(chunk) < len(chunks) - 1 or page_num < num_pages - 1:
                c.showPage()
                c.setFont("Amiri", 12)  # Reset the font for the new page

    # Save the translated PDF
    c.save()
    return translated_pdf_path

# Example usage
#pdf_path = r"C:\Users\Rayen\Downloads\data\Climate-Change.pdf"  # Replace with the path to your PDF file
#translated_pdf_path = translate_pdf(pdf_path)
#print(f"Translated PDF saved to: {translated_pdf_path}")
def translate_docx(docx_path, dest_lang='en'):
    translator = GoogleTranslator(target=dest_lang)

    # Open the original DOCX file
    doc = Document(docx_path)
    translated_doc = Document()

    # Get the name of the DOCX file
    docx_name = os.path.basename(docx_path)
    
    # Construct the translated DOCX path
    translated_docx_path = os.path.join("data", f"{os.path.splitext(docx_name)[0]}_translated.docx")

    # Iterate through each paragraph
    for para in doc.paragraphs:
        text = para.text

        # Split text into chunks of 4000 characters (adjust as needed)
        chunk_size = 4000
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

        for chunk in chunks:
            if chunk.strip():  # Only translate non-empty chunks
                translated_text = translator.translate(chunk)
                translated_doc.add_paragraph(translated_text)
            else:
                translated_doc.add_paragraph("")  # Preserve empty paragraphs

    # Save the translated DOCX file
    translated_doc.save(translated_docx_path)
    return translated_docx_path

# Example usage
#pdf_path = r"C:\Users\Rayen\Downloads\data\Climate-Change.pdf"  # Replace with the path to your PDF file
#translate_pdf(pdf_path)
docx_path = r"C:\Users\Rayen\Downloads\arabic-1333488651 (3).docx"  # Replace with the path to your DOCX file
translated_docx_path = translate_docx(docx_path)